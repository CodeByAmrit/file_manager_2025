import mysql.connector
from docx import Document
from mysql.connector import IntegrityError
from tkinter import messagebox, filedialog
from PIL import ImageTk, Image
import os, sys, io
import save_to_file as save
import pythoncom


def resource_path(relative_path):
    base_path = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)

file_path = resource_path("output.txt")

def connect():
    host, port = save.read_from_file(file_path)
    conn = mysql.connector.connect(
    host=host,
    user="filemanager",
    password="admin",
    database="bvm",
    port=port
    )
    return conn

try:
    conn = connect()
except mysql.connector.Error as err:
    messagebox.showinfo("Offline", "Your database is offline")
    save.take_inputs()
    sys.exit()

except Exception as e:
    messagebox.showwarning("Error",f"{e}")

def resource_path(relative_path):
    base_path = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)

cursor = conn.cursor(prepared=True)

def insert_into_marks1(self, srn, data):
    if not self.english_first_entry.get():
        messagebox.showwarning("Requirements First Term", "English marks* must be filled ")
        return
    if not self.hindi_first_entry.get():
        messagebox.showwarning("Requirements First Term", "Hindi marks* must be filled ")
        return
    if not self.maths_first_entry.get():
        messagebox.showwarning("Requirements First Term", "Maths marks* must be filled ")
        return
    if not self.sst_first_entry.get():
        messagebox.showwarning("Requirements First Term", "Sst marks* must be filled ")
        return
    if not self.science_first_entry.get():
        messagebox.showwarning("Requirements First Term", "Science marks* must be filled ")
        return
    if not self.computer_first_entry.get():
        messagebox.showwarning("Requirements First Term", "Computer marks* must be filled ")
        return
    if not self.drawing_first_entry.get():
        messagebox.showwarning("Requirements First Term", "Drawing marks* must be filled ")
        return
    if not self.general_first_entry.get():
        messagebox.showwarning("Requirements First Term", "General marks* must be filled ")
        return
    if not self.grand_first_entry.get():
        messagebox.showwarning("Requirements First Term", "Grand marks* must be filled ")
        return
    
    
    try:
        data = (
            srn,
            data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], 
            data[8], data[9], data[10]
        )

        query = """
            INSERT INTO marks1 (id, english1, hindi1, mathematics1, social_science1, 
                                science1, computer1, drawing1, gn1, grandTotal1, percentage1, rank1)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                english1 = VALUES(english1),
                hindi1 = VALUES(hindi1),
                mathematics1 = VALUES(mathematics1),
                social_science1 = VALUES(social_science1),
                science1 = VALUES(science1),
                computer1 = VALUES(computer1),
                drawing1 = VALUES(drawing1),
                gn1 = VALUES(gn1),
                grandTotal1 = VALUES(grandTotal1),
                percentage1 = VALUES(percentage1),
                rank1 = VALUES(rank1)
        """

        cursor.execute(query, data)
        conn.commit()
        messagebox.showinfo(
            f"Registration Completed",
            f"SRN No - {srn} Marks of First term sucessfully saved.",
        )
    except mysql.connector.IntegrityError as e:
        cursor.execute(
            "UPDATE marks1 SET english1 = ?, hindi1 = ?, mathematics1 = ?, social_science1 = ?, science1 = ?, computer1 = ?, drawing1 = ?, gn1 = ?, grandTotal1 = ?, percentage1 = ?, rank1 = ? WHERE id = ?",
            (
                data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], data[8], data[9], data[10],
                srn,
            ),
        )

        conn.commit()
        messagebox.showinfo(
            f"Database Updated",
            f"Record Updated for - {srn}",
        )
        return 0
    except Exception as e:
        messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        return 0
    
def insert_into_marks2(self, srn, data):
    if not self.english_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "English marks* must be filled ")
        return
    if not self.hindi_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "Hindi marks* must be filled ")
        return
    if not self.maths_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "Maths marks* must be filled ")
        return
    if not self.sst_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "Sst marks* must be filled ")
        return
    if not self.science_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "Science marks* must be filled ")
        return
    if not self.computer_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "Computer marks* must be filled ")
        return
    if not self.drawing_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "Drawing marks* must be filled ")
        return
    if not self.general_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "General marks* must be filled ")
        return
    if not self.grand_second_entry.get():
        messagebox.showwarning("Requirements Second Term", "Grand marks* must be filled ")
        return
    
    
    try:
        data = (
            srn,
            data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], 
            data[8], data[9], data[10]
        )

        query = """
            INSERT INTO marks2 (id, english2, hindi2, mathematics2, social_science2, 
                                science2, computer2, drawing2, gn2, grandTotal2, percentage2, rank2)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                english2 = VALUES(english2),
                hindi2 = VALUES(hindi2),
                mathematics2 = VALUES(mathematics2),
                social_science2 = VALUES(social_science2),
                science2 = VALUES(science2),
                computer2 = VALUES(computer2),
                drawing2 = VALUES(drawing2),
                gn2 = VALUES(gn2),
                grandTotal2 = VALUES(grandTotal2),
                percentage2 = VALUES(percentage2),
                rank2 = VALUES(rank2)
        """

        cursor.execute(query, data)
        conn.commit()
        messagebox.showinfo(
            f"Registration Completed",
            f"SRN No - {srn} Marks of First term sucessfully saved.",
        )
    except mysql.connector.IntegrityError as e:
        cursor.execute(
            "UPDATE marks2 SET english2 = ?, hindi2 = ?, mathematics2 = ?, social_science2 = ?, science2 = ?, computer2 = ?, drawing2 = ?, gn2 = ?, grandTotal2 = ?, percentage2 = ?, rank2 = ? WHERE id = ?",
            (
                data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], data[8], data[9], data[10],
                srn,
            ),
        )

        conn.commit()
        messagebox.showinfo(
            f"Database Updated - Second Term",
            f"Record Updated for - {srn}",
        )
        return 0
    except Exception as e:
        messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        return 0
    
def insert_into_marks3(self, srn, data):
    if not self.english_third_entry.get():
        messagebox.showwarning("Requirements 3", "English marks* must be filled ")
        return
    if not self.hindi_third_entry.get():
        messagebox.showwarning("Requirements 3", "Hindi marks* must be filled ")
        return
    if not self.maths_third_entry.get():
        messagebox.showwarning("Requirements 3", "Maths marks* must be filled ")
        return
    if not self.sst_third_entry.get():
        messagebox.showwarning("Requirements 3", "Sst marks* must be filled ")
        return
    if not self.science_third_entry.get():
        messagebox.showwarning("Requirements 3", "Science marks* must be filled ")
        return
    if not self.computer_third_entry.get():
        messagebox.showwarning("Requirements 3", "Computer marks* must be filled ")
        return
    if not self.drawing_third_entry.get():
        messagebox.showwarning("Requirements 3", "Drawing marks* must be filled ")
        return
    if not self.general_third_entry.get():
        messagebox.showwarning("Requirements 3", "General marks* must be filled ")
        return
    if not self.grand_third_entry.get():
        messagebox.showwarning("Requirements 3", "Grand marks* must be filled ")
        return
    
    
    try:
        data = (
            srn,
            data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], 
            data[8], data[9], data[10]
        )

        query = """
            INSERT INTO marks3 (id, english3, hindi3, mathematics3, social_science3, 
                                science3, computer3, drawing3, gn3, grandTotal3, percentage3, rank3)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                english3 = VALUES(english3),
                hindi3 = VALUES(hindi3),
                mathematics3 = VALUES(mathematics3),
                social_science3 = VALUES(social_science3),
                science3 = VALUES(science3),
                computer3 = VALUES(computer3),
                drawing3 = VALUES(drawing3),
                gn3 = VALUES(gn3),
                grandTotal3 = VALUES(grandTotal3),
                percentage3 = VALUES(percentage3),
                rank3 = VALUES(rank3)
        """

        cursor.execute(query, data)
        conn.commit()
        messagebox.showinfo(
            f"Registration Completed",
            f"SRN No - {srn} Marks of First term sucessfully saved.",
        )
    except mysql.connector.IntegrityError as e:
        cursor.execute(
            "UPDATE marks3 SET english3 = ?, hindi3 = ?, mathematics3 = ?, social_science3 = ?, science3 = ?, computer3 = ?, drawing3 = ?, gn3 = ?, grandTotal3 = ?, percentage3 = ?, rank3 = ? WHERE id = ?",
            (
                data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], data[8], data[9], data[10],
                srn,
            ),
        )

        conn.commit()
        messagebox.showinfo(
            f"Database Updated - Second Term",
            f"Record Updated for - {srn}",
        )
        return 0
    except Exception as e:
        messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        return 0
    
def insert_into_maximum(self, srn, data):
    if not self.english_max_entry.get():
        messagebox.showwarning("Requirements Max", "English marks* must be filled ")
        return
    if not self.hindi_max_entry.get():
        messagebox.showwarning("Requirements Max", "Hindi marks* must be filled ")
        return
    if not self.maths_max_entry.get():
        messagebox.showwarning("Requirements Max", "Maths marks* must be filled ")
        return
    if not self.sst_max_entry.get():
        messagebox.showwarning("Requirements Max", "Sst marks* must be filled ")
        return
    if not self.science_max_entry.get():
        messagebox.showwarning("Requirements Max", "Science marks* must be filled ")
        return
    if not self.computer_max_entry.get():
        messagebox.showwarning("Requirements Max", "Computer marks* must be filled ")
        return
    if not self.drawing_max_entry.get():
        messagebox.showwarning("Requirements Max", "Drawing marks* must be filled ")
        return
    if not self.general_max_entry.get():
        messagebox.showwarning("Requirements Max", "General marks* must be filled ")
        return
    if not self.grand_max_entry.get():
        messagebox.showwarning("Requirements Max", "Grand marks* must be filled ")
        return
    
    
    try:
        data = (
            srn,
            data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], 
            data[8], data[9], data[10]
        )

        query = """
            INSERT INTO maximum_marks (id, maxEng, maxHindi, maxMaths, maxSst, maxScience, 
                                    maxComp, maxDrawing, maxGn, maxGrandTotal, attendance, maxRank)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                maxEng = VALUES(maxEng),
                maxHindi = VALUES(maxHindi),
                maxMaths = VALUES(maxMaths),
                maxSst = VALUES(maxSst),
                maxScience = VALUES(maxScience),
                maxComp = VALUES(maxComp),
                maxDrawing = VALUES(maxDrawing),
                maxGn = VALUES(maxGn),
                maxGrandTotal = VALUES(maxGrandTotal),
                attendance = VALUES(attendance),
                maxRank = VALUES(maxRank)
        """

        cursor.execute(query, data)
        conn.commit()
        messagebox.showinfo(
            f"Registration Completed",
            f"SRN No - {srn} Maximum Marks sucessfully saved.",
        )
    except mysql.connector.IntegrityError as e2:
        cursor.execute(
            "UPDATE maximum_marks SET maxEng = ?, maxHindi = ?, maxMaths = ?, maxSst = ?, maxScience = ?, maxComp = ?, maxDrawing = ?, maxGn = ?, maxGrandTotal = ?, attendance = ?, maxRank = ? WHERE id = ?",
            (
                data[0], data[1], data[2], data[3], data[4], data[5], data[6], data[7], data[8], data[9], data[10],
                srn,
            ),
        )
        conn.commit()
        messagebox.showinfo(
            f"Database Updated - Second Term",
            f"Record Updated for - {srn}",
        )
        return 0
    except Exception as e:
        messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        return 0


def insert(self):
    if not self.entryofName.get():
        messagebox.showwarning("Requirements", "Name* must be filled ")
        return
    if not self.entryofFather.get():
        messagebox.showwarning("Requirements", "Father* must be filled ")
        return
    if not self.entryofMother.get():
        messagebox.showwarning("Requirements", "Mother* must be filled ")
        return
    if not self.entryofSRN_no.get():
        messagebox.showwarning("Requirements", "SRN_no* must be filled ")
        return
    if not self.entryofPEN_no.get():
        messagebox.showwarning("Requirements", "PEN_no* must be filled ")
        return
    if not self.entryofAdmission.get():
        messagebox.showwarning("Requirements", "Admission* must be filled ")
        return
    if not self.entryofSession.get():
        messagebox.showwarning("Requirements", "Session* must be filled ")
        return
    if not self.entryofRoll.get():
        messagebox.showwarning("Requirements", "Roll Number* must be filled ")
        return
    result = insert_student(
        self.entryofName.get(),
        self.entryofFather.get(),
        self.entryofMother.get(),
        self.entryofSRN_no.get(),
        self.entryofPEN_no.get(),
        self.entryofAdmission.get(),
        self.entryofClass.get(),
        self.entryofSession.get(),
        self.entryofRoll.get(),
    )
    if result == 1:
        res = self.student_photo.crop_and_save(self.entryofSRN_no.get())
        if res != 1:
            messagebox.showwarning("Image Error", f"Image Error occur - {res}")
        if res == 1:
            messagebox.showinfo(
                "Image Uploaded",
                f"Photo of student SRN NO:- {self.entryofSRN_no.get()} is Uploaded Successfully",
            )


def search_students(
    name=None, srn_no=None, pen_no=None, admission_no=None, father_name=None, clas=None, session = None, roll = None
):
    
    query = "SELECT * FROM students WHERE "
    conditions = []
    if name:
        conditions.append(f"name LIKE '%{name}%'")
    if srn_no:
        conditions.append(f"srn_no LIKE '%{srn_no}%'")
    if pen_no:
        conditions.append(f"pen_no LIKE '%{pen_no}%'")
    if admission_no:
        conditions.append(f"admission_no LIKE '%{admission_no}%'")
    if father_name:
        conditions.append(f"father_name LIKE '%{father_name}%'")
    if clas:
        conditions.append(f"class LIKE '%{clas}%'")
    if session:
        conditions.append(f"session LIKE '%{session}%'")
    if roll:
        conditions.append(f"roll LIKE '%{roll}%'")
    if conditions:
        query += " " + " AND ".join(conditions)

    try:
        
        cursor.execute(query)
        results = cursor.fetchall()
        return results
    except:
        cursor.execute("SELECT * FROM students")
        results = cursor.fetchall()
        return results
    
    
def get_first_term_marks(srn_no):
    query = f"SELECT * FROM marks1 WHERE id = {srn_no}"
    try:
        cursor.execute(query)
        results = cursor.fetchall()
        return results
    except:
        return 0
def get_second_term_marks(srn_no):
    query = f"SELECT * FROM marks2 WHERE id = {srn_no}"
    try:
        cursor.execute(query)
        results = cursor.fetchall()
        return results
    except:
        return 0
def get_third_term_marks(srn_no):
    query = f"SELECT * FROM marks3 WHERE id = {srn_no}"
    try:
        cursor.execute(query)
        results = cursor.fetchall()
        return results
    except:
        return 0
def get_max_term_marks(srn_no):
    query = f"SELECT * FROM maximum_marks WHERE id = {srn_no}"
    try:
        cursor.execute(query)
        results = cursor.fetchall()
        return results
    except:
        return 0


def insert_student(name, father_name, mother_name, srn_no, pen_no, admission_no, clas, session, roll):
    # Insert data into the "students" table
    # insert_query = "INSERT INTO students (name, father_name, mother_name, srn_no, pen_no, admission_no, class, session, roll) " \
                # #    "VALUES (%(name)s, %(father_name)s, %(mother_name)s, %(srn_no)s, %(pen_no)s, %(admission_no)s, %(class)s, %(session)s, %(roll)s)"
    # cursor.execute(insert_query, student_data)
    # connection.commit()
    
    
    
    
    
    
    
    
    
    try:
        cursor.execute(
            """
            INSERT INTO students (name, father_name, mother_name, srn_no,  pen_no, admission_no, class, session, roll)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
            (
                name.upper(),
                father_name.upper(),
                mother_name.upper(),
                srn_no.upper(),
                pen_no.upper(),
                admission_no,
                clas.upper(),
                session,
                roll,
            ),
        )
        conn.commit()
        messagebox.showinfo(
            f"Registration Completed",
            f"Student Name - {name.upper()}, SRN No - {srn_no} register sucessfully.",
        )
        return 1
    except mysql.connector.IntegrityError as e:
        messagebox.showwarning(
            f"Database Error",
            f"Record Already present for Name - {name.upper()}, SRN - {srn_no}",
        )
        return 0
    except Exception as e:
        messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        return 0


def update_student(name, father_name, mother_name, srn_no, pen_no, admission_no, clas, session, roll):
    try:
        cursor.execute(
            "UPDATE students SET name= ? , father_name= ? , mother_name= ? ,  pen_no= ? , admission_no= ? , class = ? , session = ?, roll = ? WHERE srn_no=?",
            (
                name.upper(),
                father_name.upper(),
                mother_name.upper(),
                pen_no.upper(),
                admission_no,
                clas.upper(),
                session.upper(),
                roll.upper(),
                srn_no.upper(),
            ),
        )
        conn.commit()
        messagebox.showinfo(
            f"Update Completed",
            f"Student Name - {name.upper()}, SRN No - {srn_no} Update sucessfully.",
        )
        return 1
    except mysql.connector.IntegrityError as e:
        messagebox.showwarning(
            f"Database Error",
            f"Record Already present for Name - {name.upper()}, SRN - {srn_no}",
        )
        return 0
    except Exception as e:
        messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        return 0
    
def insert_pdf_file(srn_no, pdf_path):
    try:
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()
            cursor.execute(
                """
                INSERT INTO pdf_files (srn_no, pdf_data)
                VALUES (?, ?)
            """,
                (srn_no, pdf_data),
            )
            conn.commit()
            messagebox.showinfo("Pdf Status", "Pdf File Uploaded Successfully.")
    except mysql.connector.IntegrityError as e:
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()
            cursor.execute(
                """
                UPDATE pdf_files SET pdf_data=? WHERE srn_no=?         
                """,
                (pdf_data, srn_no),
            )
            conn.commit()
            messagebox.showinfo("Pdf Status", "Pdf File Updated.")
            # print("PDF Updated")


def deleteStudent(srn):
    action = messagebox.askyesno(
        "Permission Required",
        f"You are trying to delete record of student - {srn} permanently. Are you sure ?",
    )
    if action:
        # conn = connect(resource_path("database/student_database.db"))
        # cursor = conn.cursor(prepared=True)
        try:
            cursor.execute("DELETE FROM photo WHERE id=?", (srn,))
            conn.commit()
        except Exception as e:
            messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        try:
            cursor.execute("DELETE FROM pdf_files WHERE srn_no=?", (srn,))
            conn.commit()
        except Exception as e:
            messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        try:
            cursor.execute("DELETE FROM students WHERE srn_no=?", (srn,))
            conn.commit()
            messagebox.showinfo(
                "Deleted", f"Record for Student - {srn} delete successfully."
            )
        except Exception as e:
            messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        try:
            cursor.execute("DELETE FROM marks1 WHERE id=?", (srn,))
            conn.commit()
            
        except Exception as e:
            messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        try:
            cursor.execute("DELETE FROM marks2 WHERE id=?", (srn,))
            conn.commit()
            
        except Exception as e:
            messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
        try:
            cursor.execute("DELETE FROM marks3 WHERE id=?", (srn,))
            conn.commit()
        except Exception as e:
            messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
            
        try:
            cursor.execute("DELETE FROM maximum_marks WHERE id=?", (srn,))
            conn.commit()
            
        except Exception as e:
            messagebox.showwarning(f"Database Error", f"SQLite error: {e}\n")
            
        


def retrieve_pdf_file(srn_no):
    cursor.execute("SELECT pdf_data FROM pdf_files WHERE srn_no = ?", (srn_no,))
    result = cursor.fetchone()
    if result:
        pdfSavePath = filedialog.asksaveasfilename(
            defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")]
        )
        pdf_data = result[0]
        create_pdf_from_binary(pdf_data, pdfSavePath)
        messagebox.showinfo("", "PDF Saved Successfully")
    else:
        return None


def create_pdf_from_binary(pdf_data, output_path="output.pdf"):
    with open(output_path, "wb") as pdf_file:
        pdf_file.write(pdf_data)


def getStudentsList():
    # conn = connect(resource_path("database/student_database.db"))
    # cursor = conn.cursor(prepared=True)
    cursor.execute("SELECT * FROM students")
    student = cursor.fetchall()
    return student

# Global variable to store the original image data
current_image_data = b""

def save_image():
    global current_image_data

    file_path = filedialog.asksaveasfilename(defaultextension=".jpg", filetypes=[("JPEG files", "*.jpg"), ("All files", "*.*")])

    if file_path:
        # Open the image using Pillow directly from bytes
        img = Image.frombytes("RGB", (300, 380), current_image_data)

        # Resize the image if needed
        img_resized = img.resize((250, 330), Image.LANCZOS)

        # Save the image in JPEG format
        img_resized.save(file_path, "JPEG")

def show_images_from_db(srn):
    global current_image_data

    # Assuming cursor is defined before this function
    cursor.execute("SELECT image FROM photo WHERE id=?", (srn,))
    rows = cursor.fetchall()
    if rows:
        for row in rows:
            img_bytes = row[0]
            current_image_data = img_bytes  # Store the original image data

            img = Image.frombytes("RGB", (300, 380), img_bytes)
            img_resized = img.resize((250, 330), Image.LANCZOS)
            img_tk = ImageTk.PhotoImage(img_resized)

            return img_tk
    else:
        return ""
    
def promote(srn, class_list, trigger, btn):
    if len(srn) < 1:
        messagebox.showwarning("No Student is Selected","Please Select atleast one student!")
    else:
        # conn = connect(resource_path("database/student_database.db"))
        # cursor = conn.cursor(prepared=True)
        # cursor.execute("SELECT image FROM photo WHERE id=?", (srn,),)
        # rows = cursor.fetchall()
        try:
            for id in srn:
                index = class_list.index(id[1])+trigger
                next_class = class_list[index]
                cursor.execute(
                    "UPDATE students SET class= ?  WHERE srn_no=?",
                    (
                       next_class, id[0],
                    ),
                )
                conn.commit()
            btn.search_students()
            messagebox.showinfo(
                f"Update Completed",
                f"Students Sucessfully Promoted.",
            )
        except Exception as e:
            messagebox.showerror("Database Error",f"{e}")

def replace_text_in_docx(input_docx_path, replacements):
    pythoncom.CoInitialize()
    # Load the Word document
    doc = Document(input_docx_path)

    # Iterate through all paragraphs in the document
    for para in doc.paragraphs:
        # Iterate through each replacement pair
        for old_text, new_text in replacements.items():
            # Replace the old_text with the new_text in each paragraph
            para.text = para.text.replace(old_text, new_text)

    # Iterate through all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Iterate through each replacement pair
                for old_text, new_text in replacements.items():
                    # Replace the old_text with the new_text in each cell of the table
                    cell.text = cell.text.replace(old_text, new_text)


    # Save the document to the temporary DOCX file
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    doc.save(file_path)

    messagebox.showinfo('Certificate Generated', f'Certificate generated successfully!')
    pythoncom.CoUninitialize()


if __name__ == "__main__":
    pass
            
